from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def apply_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for idx, line in enumerate(lines):
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if idx == 1 and all(set(cell) <= {":", "-"} for cell in cells):
            continue
        rows.append(cells)
    return rows


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(10.5)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Inter"


def build_docx(source_path: Path, output_path: Path) -> None:
    doc = Document()
    set_base_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    lines = source_path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            i += 1
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.+?)\)$", stripped)
        if image_match:
            image_path = Path(image_match.group(2)).expanduser()
            if not image_path.is_absolute():
                image_path = (source_path.parent / image_path).resolve()
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.5))
                if image_match.group(1):
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    apply_runs(caption, image_match.group(1))
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            paragraph = doc.add_paragraph()
            paragraph.style = doc.styles["No Spacing"]
            run = paragraph.add_run("\n".join(block))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        cell_paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        apply_runs(cell_paragraph, cell)
                        if r_idx == 0:
                            for run in cell_paragraph.runs:
                                run.bold = True
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:])
            run.bold = True
            run.font.name = "Inter"
            run.font.size = Pt(16)
            i += 1
            continue

        if stripped.startswith("## "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[3:])
            run.font.name = "Inter"
            run.font.size = Pt(11)
            i += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            apply_runs(paragraph, stripped[4:])
            i += 1
            continue

        if stripped.startswith("##") or stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            style = "Heading 1" if level == 1 else "Heading 2"
            paragraph = doc.add_paragraph(style=style)
            apply_runs(paragraph, text)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item_text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                paragraph = doc.add_paragraph(style="List Number")
                apply_runs(paragraph, item_text)
                i += 1
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                item_text = lines[i].strip()[2:]
                paragraph = doc.add_paragraph(style="List Bullet")
                apply_runs(paragraph, item_text)
                i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            if next_line == "---" or next_line.startswith("#") or next_line.startswith("|") or next_line.startswith("```") or next_line.startswith("- ") or re.match(r"^\d+\.\s+", next_line):
                break
            paragraph_lines.append(next_line)
            i += 1

        paragraph = doc.add_paragraph()
        apply_runs(paragraph, " ".join(paragraph_lines))

    doc.save(str(output_path))


def main() -> None:
    if len(sys.argv) >= 3:
        source = Path(sys.argv[1])
        output = Path(sys.argv[2])
    else:
        source = Path("CM3070_full_report_draft.md")
        output = Path("CM3070_full_report_draft.docx")
    build_docx(source, output)
    print(f"Created DOCX at {output}")


if __name__ == "__main__":
    main()
